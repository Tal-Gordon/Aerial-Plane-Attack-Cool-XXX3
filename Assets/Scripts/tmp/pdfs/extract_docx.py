from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph


def iter_blocks(parent):
    for child in parent.element.body.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, parent)
        elif child.tag.endswith("}tbl"):
            yield Table(child, parent)


def extract(path: Path) -> str:
    doc = Document(path)
    lines: list[str] = []
    props = doc.core_properties
    lines.extend(
        [
            f"# {path.name}",
            "",
            f"- Title: {props.title or ''}",
            f"- Subject: {props.subject or ''}",
            f"- Author: {props.author or ''}",
            f"- Sections: {len(doc.sections)}",
            f"- Inline images: {len(doc.inline_shapes)}",
            "",
        ]
    )

    for block in iter_blocks(doc):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if not text:
                continue
            style = block.style.name if block.style else ""
            prefix = ""
            lowered = style.lower()
            if lowered.startswith("heading"):
                try:
                    level = int(lowered.split()[-1])
                except ValueError:
                    level = 2
                prefix = "#" * min(level + 1, 6) + " "
            elif "title" in lowered:
                prefix = "## "
            elif "list" in lowered:
                prefix = "- "
            lines.append(f"{prefix}{text}")
            lines.append("")
        else:
            rows = []
            for row in block.rows:
                cells = [" ".join(cell.text.split()) for cell in row.cells]
                rows.append(cells)
            if not rows:
                continue
            width = max(len(row) for row in rows)
            rows = [row + [""] * (width - len(row)) for row in rows]
            lines.append("| " + " | ".join(rows[0]) + " |")
            lines.append("| " + " | ".join(["---"] * width) + " |")
            for row in rows[1:]:
                lines.append("| " + " | ".join(row) + " |")
            lines.append("")

    for section_index, section in enumerate(doc.sections, start=1):
        header = " ".join(p.text.strip() for p in section.header.paragraphs if p.text.strip())
        footer = " ".join(p.text.strip() for p in section.footer.paragraphs if p.text.strip())
        if header or footer:
            lines.append(f"## Section {section_index} furniture")
            if header:
                lines.append(f"- Header: {header}")
            if footer:
                lines.append(f"- Footer: {footer}")
            lines.append("")

    return "\n".join(lines)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("input", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    args.output.write_text(extract(args.input), encoding="utf-8")


if __name__ == "__main__":
    main()
